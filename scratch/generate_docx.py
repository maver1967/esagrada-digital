import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_callout(doc, text, title="NOTA IMPORTANTE", border_hex="172A52", bg_hex="F0F4F8"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="36" w:space="0" w:color="{border_hex}"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"📌 {title}\n")
    run_t.bold = True
    run_t.font.name = "Arial"
    run_t.font.size = Pt(10)
    run_t.font.color.rgb = RGBColor(0x17, 0x2A, 0x52)
    
    run_b = p.add_run(text)
    run_b.font.name = "Arial"
    run_b.font.size = Pt(9.5)
    run_b.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)

def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "272C34")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="12" w:space="0" w:color="3E4451"/><w:top w:val="single" w:sz="12" w:space="0" w:color="3E4451"/><w:right w:val="single" w:sz="12" w:space="0" w:color="3E4451"/><w:bottom w:val="single" w:sz="12" w:space="0" w:color="3E4451"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xAB, 0xB2, 0xBF)
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)

def create_manual_doc(output_path):
    doc = docx.Document()
    
    # Page Setup
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Title Header Box
    tbl_hdr = doc.add_table(rows=1, cols=1)
    tbl_hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_hdr = tbl_hdr.cell(0, 0)
    set_cell_background(c_hdr, "172A52")
    set_cell_margins(c_hdr, top=200, bottom=200, left=240, right=240)
    
    p_title = c_hdr.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after = Pt(4)
    
    r_sub = p_title.add_run("ESCOLA PRÉ-UNIVERSITÁRIA SAGRADA FAMÍLIA · MAXIXE\nDEPARTAMENTO DE INFRAESTRUTURA E TECNOLOGIAS DE INFORMAÇÃO\n\n")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(9)
    r_sub.bold = True
    r_sub.font.color.rgb = RGBColor(0xD0, 0xDC, 0xF0)
    
    r_main = p_title.add_run("GUIA DE INSTALAÇÃO DEFINITIVA DA PLATAFORMA DIGITAL\nNAS SYNOLOGY DS923+ (DSM 7.x)")
    r_main.font.name = "Arial"
    r_main.font.size = Pt(16)
    r_main.bold = True
    r_main.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(8)
    p_meta.paragraph_format.space_after = Pt(16)
    r_m = p_meta.add_run("Aplicação: ESAGRADA Gestor Modern v4.2.2  |  Data: Agosto 2026  |  Servidor Alvo: Synology DS923+")
    r_m.font.name = "Arial"
    r_m.font.size = Pt(9)
    r_m.font.italic = True
    r_m.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    # Section 1
    h1 = doc.add_heading("1. Visão Geral da Arquitetura e Requisitos", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x17, 0x2A, 0x52)
    h1.runs[0].font.name = "Arial"
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run("A plataforma ")
    r.font.name = "Arial"; r.font.size = Pt(10)
    r_b = p.add_run("ESAGRADA Gestor Modern")
    r_b.bold = True; r_b.font.name = "Arial"; r_b.font.size = Pt(10)
    r_2 = p.add_run(" é uma aplicação Web de elevada otimização baseada na arquitetura ")
    r_2.font.name = "Arial"; r_2.font.size = Pt(10)
    r_spa = p.add_run("Single Page Application (SPA) / Progressive Web App (PWA)")
    r_spa.bold = True; r_spa.font.name = "Arial"; r_spa.font.size = Pt(10)
    r_3 = p.add_run(". O processamento de dados e o armazenamento de turmas, cadernetas e pautas ocorrem localmente no navegador via IndexedDB, enquanto o servidor NAS fornece a distribuição rápida de ficheiros estáticos e suporte ao Service Worker.")
    r_3.font.name = "Arial"; r_3.font.size = Pt(10)
    
    # Key Specs Bullet List
    bullets = [
        ("Arquitetura:", " Client-Side SPA compilada com Vite (HTML5, CSS3, JS ES6+)."),
        ("Armazenamento de Dados:", " IndexedDB e LocalStorage local no navegador."),
        ("Requisito Crítico (PWA):", " Ligação HTTPS (porta 443 ou 8443) com Certificado SSL ativo (obrigatório para instalação no telemóvel e funcionamento offline)."),
        ("Servidor de Produção:", " Synology DS923+ a correr DSM 7.0+ com pacote Web Station ou Container Manager.")
    ]
    for bold_prefix, text in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_before = Pt(1)
        bp.paragraph_format.space_after = Pt(3)
        rb = bp.add_run(bold_prefix)
        rb.bold = True; rb.font.name = "Arial"; rb.font.size = Pt(9.5)
        rt = bp.add_run(text)
        rt.font.name = "Arial"; rt.font.size = Pt(9.5)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    
    # Section 2
    h2 = doc.add_heading("2. Método Recomendado: Instalação via Synology Web Station", level=1)
    h2.runs[0].font.color.rgb = RGBColor(0x17, 0x2A, 0x52)
    h2.runs[0].font.name = "Arial"
    
    add_callout(doc, 
                "Este é o método mais eficiente e leve para o Synology DS923+. Integra diretamente a aplicação estática no servidor Nginx nativo do DSM, consumindo recursos mínimos de RAM e CPU e integrando-se nativamente com a gestão de certificados SSL do Synology.",
                title="VANTAGENS DO MÉTODO WEB STATION")

    steps = [
        ("Etapa 1: Compilação dos Ficheiros de Produção (Build)", [
            "1. No computador de desenvolvimento, abra o terminal na pasta do projeto esagrada-gestor-modern.",
            "2. Execute o comando: npm run build",
            "3. O processo gera a pasta 'dist/' contendo index.html, app.js, style.css, manifest.webmanifest e recursos."
        ]),
        ("Etapa 2: Preparação do Servidor e Pasta no NAS", [
            "1. Aceda ao DSM do Synology no navegador.",
            "2. Abra o File Station e navegue até à pasta partilhada 'web/'.",
            "3. Crie a subpasta: /web/esagrada/",
            "4. Copie TODO o conteúdo da pasta 'dist/' para /web/esagrada/.",
            "5. Verifique se o grupo 'http' tem permissões de Leitura na pasta /web/esagrada."
        ]),
        ("Etapa 3: Instalação dos Pacotes no DSM", [
            "1. No DSM, abra o Centro de Pacotes (Package Center).",
            "2. Procure por 'Web Station' e clique em Instalar.",
            "3. Garanta que o servidor HTTP Nginx está instalado e ativo."
        ]),
        ("Etapa 4: Configuração do Serviço Web e Portal Web", [
            "1. Abra o pacote Web Station no menu principal do DSM.",
            "2. Vá a 'Serviço Web' > Criar > Selecione 'Website Estático'.",
            "3. Nome: esagrada-app  |  Pasta Raiz: /web/esagrada",
            "4. Vá a 'Portal Web' > Criar > 'Portal de Serviço Web'.",
            "5. Escolha o serviço 'esagrada-app'. Defina a porta HTTPS personalizada (ex: 8443) ou Host FQDN (ex: gestor.esagrada.local)."
        ]),
        ("Etapa 5: Certificado SSL e HTTPS (Obrigatório para PWA)", [
            "1. Abra o Painel de Controlo > Segurança > Certificado.",
            "2. Adicione um Certificado (Let's Encrypt para domínio público/DDNS ou Certificado Local).",
            "3. Clique em 'Definições' e associe o certificado ao serviço 'esagrada-app'."
        ])
    ]
    
    for title, substeps in steps:
        h3 = doc.add_heading(title, level=2)
        h3.runs[0].font.color.rgb = RGBColor(0x2B, 0x4C, 0x7E)
        h3.runs[0].font.name = "Arial"
        for s in substeps:
            p_sub = doc.add_paragraph()
            p_sub.paragraph_format.space_before = Pt(1)
            p_sub.paragraph_format.space_after = Pt(2)
            p_sub.paragraph_format.left_indent = Inches(0.2)
            r_s = p_sub.add_run(s)
            r_s.font.name = "Arial"; r_s.font.size = Pt(9.5)
            
    # Section 3: Docker Method
    h3_docker = doc.add_heading("3. Método Alternativo: Implantação via Container Manager (Docker)", level=1)
    h3_docker.runs[0].font.color.rgb = RGBColor(0x17, 0x2A, 0x52)
    h3_docker.runs[0].font.name = "Arial"
    
    p_dock = doc.add_paragraph()
    p_dock.paragraph_format.space_after = Pt(6)
    r_d = p_dock.add_run("Caso a equipa prefira isolar a aplicação num contentor Docker com servidor Nginx dedicado, utilize o Container Manager do DSM 7:")
    r_d.font.name = "Arial"; r_d.font.size = Pt(10)
    
    p_compose_title = doc.add_paragraph()
    p_compose_title.paragraph_format.space_before = Pt(4)
    p_compose_title.paragraph_format.space_after = Pt(2)
    r_c_t = p_compose_title.add_run("Ficheiro: docker-compose.yml")
    r_c_t.bold = True; r_c_t.font.name = "Arial"; r_c_t.font.size = Pt(10); r_c_t.font.color.rgb = RGBColor(0x2B, 0x4C, 0x7E)
    
    code_docker = """version: '3.8'

services:
  esagrada-web:
    image: nginx:alpine
    container_name: esagrada_pwa
    restart: always
    ports:
      - "8080:80"
    volumes:
      - /volume1/web/esagrada:/usr/share/nginx/html:ro
      - /volume1/docker/esagrada/nginx.conf:/etc/nginx/conf.d/default.conf:ro"""
    add_code_block(doc, code_docker)
    
    p_nginx_title = doc.add_paragraph()
    p_nginx_title.paragraph_format.space_before = Pt(4)
    p_nginx_title.paragraph_format.space_after = Pt(2)
    r_n_t = p_nginx_title.add_run("Ficheiro: /volume1/docker/esagrada/nginx.conf")
    r_n_t.bold = True; r_n_t.font.name = "Arial"; r_n_t.font.size = Pt(10); r_n_t.font.color.rgb = RGBColor(0x2B, 0x4C, 0x7E)
    
    code_nginx = """server {
    listen 80;
    server_name _;
    root /usr/share/nginx/html;
    index index.html;

    location / {
        try_files $uri $uri/ /index.html;
    }

    location ~* \\.(js|css|png|jpg|jpeg|gif|ico|svg|woff2)$ {
        expires 30d;
        add_header Cache-Control "public, no-transform";
    }

    location ~* (sw\\.js|manifest\\.webmanifest)$ {
        expires -1;
        add_header Cache-Control "no-store, no-cache, must-revalidate";
    }
}"""
    add_code_block(doc, code_nginx)

    # Section 4: Network & Updates
    h4 = doc.add_heading("4. Configuração de Rede Local e Acesso", level=1)
    h4.runs[0].font.color.rgb = RGBColor(0x17, 0x2A, 0x52)
    h4.runs[0].font.name = "Arial"
    
    net_items = [
        ("IP Estático no NAS:", " Defina um IP fixo em Painel de Controlo > Rede > Interface (ex: 192.168.1.100)."),
        ("DNS Local:", " Crie um registo A no router/DNS interno apontando gestor.esagrada.local para o IP do NAS."),
        ("Acesso dos Utilizadores:", " Os professores acedem via https://gestor.esagrada.local ou https://192.168.1.100:8443.")
    ]
    for b_prefix, txt in net_items:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_before = Pt(1)
        bp.paragraph_format.space_after = Pt(3)
        rb = bp.add_run(b_prefix)
        rb.bold = True; rb.font.name = "Arial"; rb.font.size = Pt(9.5)
        rt = bp.add_run(txt)
        rt.font.name = "Arial"; rt.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    h5 = doc.add_heading("5. Procedimento de Atualização (Workflow de Manutenção)", level=1)
    h5.runs[0].font.color.rgb = RGBColor(0x17, 0x2A, 0x52)
    h5.runs[0].font.name = "Arial"
    
    upd_items = [
        "1. No computador do desenvolvedor: Executar npm run build.",
        "2. No File Station do Synology: Substituir o conteúdo de /web/esagrada/ pelos novos ficheiros da pasta dist/.",
        "3. Nos telemóveis/PCs dos professores: Clicar no botão '🔄 atualizar' no rodapé da barra lateral da app para purgar a cache do Service Worker instantaneamente."
    ]
    for u_txt in upd_items:
        p_u = doc.add_paragraph()
        p_u.paragraph_format.space_before = Pt(1)
        p_u.paragraph_format.space_after = Pt(3)
        r_u = p_u.add_run(u_txt)
        r_u.font.name = "Arial"; r_u.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Section 6: Checklist Table
    h6 = doc.add_heading("6. Lista de Verificação Final (Checklist de TI)", level=1)
    h6.runs[0].font.color.rgb = RGBColor(0x17, 0x2A, 0x52)
    h6.runs[0].font.name = "Arial"
    
    table_data = [
        ("Item de Verificação", "Estado Esperado", "Resultado"),
        ("Ficheiros no NAS", "Conteúdo da pasta dist/ em /web/esagrada/", "✅ OK"),
        ("Permissões de Ficheiro", "Utilizador 'http' com permissão de leitura", "✅ OK"),
        ("Serviço HTTPS", "Acesso ativo em HTTPS com certificado SSL", "✅ OK"),
        ("Instalabilidade PWA", "Banner de instalação visível em dispositivos móveis", "✅ OK"),
        ("Modo Offline", "App abre normalmente mesmo sem rede após 1º acesso", "✅ OK"),
        ("Atualização de Dados", "Botão '🔄 atualizar' funcional no rodapé", "✅ OK")
    ]
    
    tbl_chk = doc.add_table(rows=len(table_data), cols=3)
    tbl_chk.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header styling
    hdr_cells = tbl_chk.rows[0].cells
    hdr_cells[0].width = Inches(2.2)
    hdr_cells[1].width = Inches(3.4)
    hdr_cells[2].width = Inches(1.0)
    for i, title in enumerate(table_data[0]):
        cell = hdr_cells[i]
        set_cell_background(cell, "172A52")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(title)
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    # Body rows styling
    for r_idx in range(1, len(table_data)):
        row_cells = tbl_chk.rows[r_idx].cells
        bg_color = "F9FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx in range(3):
            cell = row_cells[c_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            
            # cell borders
            tcPr = cell._element.get_or_add_tcPr()
            borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="0" w:color="E0E4EC"/><w:top w:val="none"/><w:left w:val="none"/><w:right w:val="none"/></w:tcBorders>')
            tcPr.append(borders)
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 2 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(table_data[r_idx][c_idx])
            r.font.name = "Arial"
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            if c_idx == 0:
                r.bold = True

    doc.save(output_path)
    print(f"Document saved successfully at: {output_path}")

if __name__ == "__main__":
    create_manual_doc("/Users/roberto67/Library/CloudStorage/Dropbox/Moçambique/Maxixe/Pré-universitária/Pré-Univeritária/esagrada-gestor-modern/Guia_Instalacao_NAS_Synology_DS923_ESAGRADA.docx")
